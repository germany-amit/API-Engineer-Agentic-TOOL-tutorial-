import streamlit as st
import json, re
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import PyPDF2   # 👈 added

st.title("📑 API Advisor from RFP")

# Upload RFP file
uploaded_file = st.file_uploader("Upload RFP (txt/pdf)", type=["txt", "pdf"])

rfp_text = ""
if uploaded_file:
    if uploaded_file.type == "application/pdf":   # 👈 handle PDF
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            rfp_text += page.extract_text() or ""
    else:  # TXT
        rfp_text = uploaded_file.read().decode("utf-8", errors="ignore")

    st.text_area("RFP Content", rfp_text, height=200)

# --- Configurable Options ---
st.subheader("1️⃣ API Type Selection")
auto_api = "REST"
if re.search(r"graphql", rfp_text, re.I):
    auto_api = "GraphQL"
elif re.search(r"real[- ]?time|grpc", rfp_text, re.I):
    auto_api = "gRPC"

api_type = st.radio("Select API type (auto-detected: %s)" % auto_api,
                   ["REST", "GraphQL", "gRPC"],
                   index=["REST","GraphQL","gRPC"].index(auto_api))

# --- Validation ---
st.subheader("2️⃣ Validation Checklist")
checks = {
    "Endpoints": r"endpoint",
    "Authentication": r"auth|login|oauth|token",
    "Data Model": r"schema|model|database",
    "Performance": r"scalability|latency|throughput"
}
missing = []
for label, pattern in checks.items():
    if not re.search(pattern, rfp_text, re.I):
        missing.append(label)

if missing:
    st.warning("Missing details in RFP: " + ", ".join(missing))
else:
    st.success("All key details seem present.")

# --- Compliance Hints ---
st.subheader("3️⃣ Compliance Hints")
compliance_flags = []
for keyword in ["personal data", "pii", "gdpr", "hipaa", "consent"]:
    if re.search(keyword, rfp_text, re.I):
        compliance_flags.append(keyword)

if compliance_flags:
    st.error("⚠️ Potential compliance requirements: " + ", ".join(compliance_flags))
else:
    st.info("No explicit GDPR/PII terms found.")

# --- Draft Spec ---
st.subheader("4️⃣ Draft API Spec")
spec = {
    "apiType": api_type,
    "endpoints": ["/sample"],
    "security": "OAuth2 (placeholder)",
    "dataModel": {"id": "integer", "name": "string"}
}
st.json(spec)

# --- Export ---
if st.button("💾 Export as JSON"):
    st.download_button("Download JSON", json.dumps(spec, indent=2), file_name="api_spec.json")

if st.button("📄 Export as PDF"):
    pdf = SimpleDocTemplate("spec.pdf")
    styles = getSampleStyleSheet()
    pdf.build([Paragraph(json.dumps(spec, indent=2), styles['Normal'])])
    with open("spec.pdf", "rb") as f:
        st.download_button("Download PDF", f, file_name="api_spec.pdf")

if st.button("📝 Export as Word"):
    doc = Document()
    doc.add_heading("API Spec", 0)
    doc.add_paragraph(json.dumps(spec, indent=2))
    doc.save("spec.docx")
    with open("spec.docx", "rb") as f:
        st.download_button("Download Word", f, file_name="api_spec.docx")

with st.container():
    st.markdown(
        """
        <div style="background-color:#f0f4f8; padding:20px; border-radius:12px;">
        <h2>🔑 Why API Engineers & Solution Architects Should Use This App</h2>

        ### Unique Features
        - **RFP Parsing** – Quickly extracts use cases, constraints, and assumptions from lengthy RFP documents.  
        - **Configurable API Options** – Lets you experiment with REST, GraphQL, or gRPC without manual rework.  
        - **Validation Checklist** – Flags missing details such as endpoints, authentication, data model, and performance expectations.  
        - **Compliance Hints** – Provides GDPR/PII reminders so security is not overlooked early.  
        - **Collaboration Ready** – Exportable spec in TXT/PDF for sharing with clients and teams.  

        ### USP (Unique Selling Points)
        - **Bridges the gap** between vague RFP language and concrete API design steps.  
        - **Lightweight & fast** – runs fully on free Streamlit, no infra or license cost.  
        - **Standardizes discovery phase** – ensures all engineers start with the same checklist and compliance mindset.  

        ### UVP (Unique Value Proposition)
        This app helps **API Engineers in Fortune 500 IT MNCs** move from *RFP text → actionable API blueprint* in minutes.  
        It reduces ambiguity, enforces consistency, and gives engineers a **head start on design decisions** —  
        saving **days of manual requirement clarification** in large enterprise projects.  
        </div>
        """,
        unsafe_allow_html=True
    )
